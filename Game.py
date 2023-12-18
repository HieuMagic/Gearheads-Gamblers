import pygame, sys, smtplib, random, re, time, Aready, TakePhotos, os, FaceID, random, datetime
from utils import load_image, load_images, get_information
from docx import Document
from pygame import mixer
from pygame.locals import *
from player import Player
from buff import Buff
from car import Car
from menu import *

class TextInput:
    def __init__(self, name, screen, font, position, color, width, height):
        self.text = name  
        self.screen = screen  
        self.font = font  
        self.position = position  
        self.color = color
        self.width = width
        self.height = height
        self.cursor_visible = True
        self.cursor_timer = 0
        self.cursor_switch_interval = 400
        self.draw()
    def draw(self):
        self.username_f = self.font.render(f'{self.text}', True, (0, 0, 0))
        self.name_box = self.username_f.get_rect(midleft=self.position)
        self.rect = pygame.draw.rect(self.screen, self.color, (self.position[0] - 5, self.position[1] - 15, self.width, self.height))
        self.screen.blit(self.username_f, self.name_box)
    def update_text(self, new_text):
        self.text = new_text
    def update_cursor(self):
        current_time = pygame.time.get_ticks()
        if current_time - self.cursor_timer > self.cursor_switch_interval:
            self.cursor_visible = not self.cursor_visible
            self.cursor_timer = current_time
        if self.cursor_visible:
            pygame.draw.line(self.screen, (0, 0, 0), (self.name_box.right + 2, self.name_box.top + 2), (self.name_box.right + 2, self.name_box.bottom - 2), 2)

class Button2:
    def __init__(self, name, screen, position, width, height):
        self.name = name
        self.screen = screen
        self.position = position
        self.width = width
        self.height = height
        self.rect = pygame.Rect(0, 0, self.width, self.height)
        self.rect.center = self.position
    def draw(self, border_radius, color):
        pygame.draw.rect(self.screen, color, self.rect, border_radius = border_radius)
        self.button_rect = self.name.get_rect(center=self.rect.center)
        self.screen.blit(self.name, self.button_rect.topleft)

class Gmail:
    def is_gmail_address(self, email):
        self.pattern = r'^[a-zA-Z0-9_.+-]+@gmail\.com$'
        return re.match(self.pattern, email) is not None
    def send_email(self, receiver_email):
        try:
            self.sender_email = "gearheadngambler@gmail.com"     
            self.password = "aadm cprn zipr fhpw"
            self.verification_code = ''.join(random.choice('0123456789') for i in range(6))
            self.session = smtplib.SMTP('smtp.gmail.com', 587)
            self.session.starttls()
            self.session.login(self.sender_email, self.password)
            self.subject = 'Subject: XÁC NHẬN ĐĂNG KÝ TÀI KHOẢN GEARHEADS&GAMBLERS!'
            self.content = f'Mã xác nhận đăng ký của bạn là {self.verification_code}.'
            self.message = f'{self.subject}\n\n{self.content}'
            self.message = self.message.encode("utf-8")
            self.session.sendmail(self.sender_email, receiver_email, self.message)
            print('Mail sent successfully')
            return True
        except Exception as e:
            print(f"Lỗi khi gửi email: {str(e)}")
            return False

class Box:
    def __init__(self, text, screen, font, position, color, width, height):
        self.text = text
        self.screen = screen
        self.font = font
        self.position = position
        self.color = color
        self.width = width
        self.height = height
        self.rect = pygame.Rect(0, 0, self.width, self.height)
        self.rect.center = self.position
    def draw(self):
        self.setup_f = self.font.render(f'{self.text}', True, (0, 0, 0))
        pygame.draw.rect(self.screen, self.color, self.rect, border_radius = 10)
        self.box_rect = self.setup_f.get_rect(center=self.rect.center)
        self.screen.blit(self.setup_f, self.box_rect.topleft)
    def update_text(self, new_text):
        self.text = new_text 

class Login_Game:
    def __init__(self):
        pygame.init()
        self.ID_Login =[0]
        self.screenWidth = 1268
        self.screenHeight = 746
        self.screen = pygame.display.set_mode((self.screenWidth, self.screenHeight))
        self.Gamename = pygame.display.set_caption('GEARHEADS&GAMBLERS')
        self.background = pygame.image.load(r'Picture\background.png').convert_alpha()
        self.logo = pygame.image.load(r'Picture/LOGO.png')
        self.logo = pygame.display.set_icon(self.logo)
        self.background = pygame.transform.scale(self.background, (self.screenWidth, self.screenHeight))
        self.Face_ID_1 = pygame.image.load(r'Picture\Face_ID_1.png')
        self.Face_ID_1 = pygame.transform.scale(self.Face_ID_1, (90, 90))
        self.Face_ID_1_rect = self.Face_ID_1.get_rect()
        self.Face_ID_1_rect.topleft = (self.screenWidth // 3 + 310, self.screenHeight // 3 + 265)
        self.Face_ID_2 = pygame.image.load(r'Picture\Face_ID_2.png')
        self.Face_ID_2 = pygame.transform.scale(self.Face_ID_2, (85, 85))
        self.Game_font = pygame.font.Font(r'Font\seguili.ttf', 20)
        self.Game_font_setup = pygame.font.Font(r'Font\segoeui.ttf', 30)
        self.Game_font_button_send = pygame.font.Font(r'Font\ARIAL.TTF', 15)
        self.Game_font_button = pygame.font.Font(r'Font\ARIAL.TTF', 20)
        self.name = "Username"
        self.pas = "Password"
        self.conpas = "Confirm password"
        self.email = "Email"
        self.verify = 'Verity'
        self.announcement = ''
        self.mode = 'LOGIN'
        self.check = ''
        self.setup = ''
        self.docx_filename = "Accounts.docx"
        self.gmail = Gmail()
        self.money = 500
        self.match = 0
        self.win = 0
        self.lose = 0
        self.cursor_click = [False, False, False, False, False]
        self.verify_check = False
        self.gmail_codecheck = False
        self.button_click_send = False
        self.Face_ID_1_click = False
        self.Clock = pygame.time.Clock()
        self.username_rect = TextInput(self.name, self.screen, self.Game_font, (self.screenWidth // 3, self.screenHeight // 3), (255, 255, 255), 400, 30)
        self.password_rect = TextInput(self.pas, self.screen, self.Game_font, (self.screenWidth // 3, self.screenHeight // 3 + 50), (255, 255, 255), 400, 30)
        self.conpasword_rect = TextInput(self.conpas, self.screen, self.Game_font, (self.screenWidth // 3, self.screenHeight // 3 + 100), (255, 255, 255), 400, 30)
        self.email_rect = TextInput(self.email, self.screen, self.Game_font, (self.screenWidth // 3, self.screenHeight // 3 + 150), (255, 255, 255), 400, 30)
        self.verify_rect = TextInput(self.verify, self.screen, self.Game_font, (self.screenWidth // 3 + 44 , self.screenHeight // 3 + 185), (255, 255, 255), 300, 30)
        self.announcement_rect = TextInput(self.announcement, self.screen, self.Game_font, (self.screenWidth // 3 , self.screenHeight // 3 + 235), (220, 240, 230), 400, 30)
        self.Setup = Box(self.setup, self.screen, self.Game_font_setup, (self.screenWidth // 3 + 200, self.screenHeight // 3 - 75), (0, 128, 0), 500, 50)
        self.Login_button = Button2(self.Game_font_button.render('LOGIN', True, (0, 0, 0)), self.screen, (self.screenWidth // 3 + 145, self.screenHeight // 3 + 285), 275, 30)
        self.SignUp_button = Button2(self.Game_font_button.render('SIGN UP', True, (0, 0, 0)), self.screen, (self.screenWidth // 3 + 145, self.screenHeight // 3 + 335), 275, 30)
        self.send_rect = Button2(self.Game_font_button_send.render('Send', True, (0, 0, 0)), self.screen, (self.screenWidth // 3 + 365, self.screenHeight // 3 + 185), 50, 30)

    def get_ID(self):
        return self.ID_Login[0]
    
    def save_account(self, username, password, money, match, win, lose, win_money, lose_money):
        try:
            if os.path.exists(self.docx_filename):
                doc = Document(self.docx_filename)
            else:
                doc = Document()
                table = doc.add_table(rows=1, cols=7)
                table.cell(0, 0).text = "ID"
                table.cell(0, 1).text = "Tên đăng nhập"
                table.cell(0, 2).text = "Mật khẩu"
                table.cell(0, 3).text = "Số tiền"
                table.cell(0, 4).text = "Số trận đã chơi"
                table.cell(0, 5).text = "Số trận thắng"
                table.cell(0, 6).text = "Số trận thua"
                table.cell(0, 7).text = "Tổng số tiền thắng cược"
                table.cell(0, 8).text = "Tổng số tiền đã thua"
            if not doc.tables:
                doc.add_table(rows=1, cols=7)
                table = doc.tables[0]
                table.cell(0, 0).text = "ID"
                table.cell(0, 1).text = "Tên đăng nhập"
                table.cell(0, 2).text = "Mật khẩu"
                table.cell(0, 3).text = "Số tiền"
                table.cell(0, 4).text = "Số trận đã chơi"
                table.cell(0, 5).text = "Số trận thắng"
                table.cell(0, 6).text = "Số trận thua"
                table.cell(0, 7).text = "Tổng số tiền thắng cược"
                table.cell(0, 8).text = "Tổng số tiền đã thua"
            for table in doc.tables:
                for row in table.rows:
                    if row.cells[1].text == username:
                        if row.cells[3].text == '0' and row.cells[4].text == '0' and row.cells[5].text == '0' and row.cells[6].text == '0':
                            print(f"Người dùng '{username}' đã tồn tại. Không thể đăng ký.")
                            return False
                        else:
                            row_cells = row.cells
                            row_cells[2].text = password
                            row_cells[3].text = str(money)
                            row_cells[4].text = str(match)
                            row_cells[5].text = str(win)
                            row_cells[6].text = str(lose)
                            row_cells[7].text = str(win_money)
                            row_cells[8].text = str(lose_money)
                            doc.save(self.docx_filename)
                            return True
            if len(doc.tables) > 0 and len(doc.tables[0].rows) > 1:
                last_id = int(doc.tables[0].cell(len(doc.tables[0].rows) - 1, 0).text)
            else:
                last_id = 0
            new_id = last_id + 1
            new_id_str = f"{new_id:03d}" 
            table = doc.tables[0] 
            row_cells = table.add_row().cells
            row_cells[0].text = new_id_str
            row_cells[1].text = username
            row_cells[2].text = password 
            row_cells[3].text = str(money)
            row_cells[4].text = str(match)
            row_cells[5].text = str(win)
            row_cells[6].text = str(lose)
            row_cells[7].text = str(win_money)
            row_cells[8].text = str(lose_money)
            print('Tài khoản đã được lưu!')
            doc.save(self.docx_filename)
            return True
        except Exception as e:
            print(f"Lỗi khi lưu tài khoản vào tệp Word: {str(e)}")
            return False

    def Count_id(self):
        doc = Document(self.docx_filename)
        if len(doc.tables) > 0 and len(doc.tables[0].rows) > 1:
            last_id = int(doc.tables[0].cell(len(doc.tables[0].rows) - 1, 0).text)
        else:
            last_id = 0
        id_count = last_id
        id_count_str = f"{id_count:03d}"
        return id_count_str
    def check_login(self, username, password):
        try:
            doc = Document(self.docx_filename)
            for table in doc.tables:
                for row in table.rows:
                    stored_username = row.cells[1].text
                    stored_password = row.cells[2].text
                    if username == stored_username and password == stored_password:
                        return True
            return False
        except Exception as e:
            print(f"Lỗi khi kiểm tra thông tin đăng nhập từ tệp Word: {str(e)}")
            return False  
        
    def SignUp(self):
        if self.mode == 'LOGIN':
            self.username_rect.update_text('Username')
            self.password_rect.update_text('Password')
            self.conpasword_rect.update_text('Confirm password')
            self.email_rect.update_text('Email')
            self.announcement = ''
            self.mode = 'SIGN UP'
            self.check = 'SIGN UP'
        elif self.mode == 'SIGN UP':
            self.check = 'SIGN UP'
            self.samename = 0
            if self.username_rect.text == "Username" or self.username_rect.text == '':
                self.announcement_rect.update_text("Please enter your username.")
            elif self.password_rect.text == "Password" or self.conpasword_rect.text == '':
                self.announcement_rect.update_text("Please enter your password.")
            elif self.conpasword_rect.text == "Confirm password" or self.password_rect.text == '':
                self.announcement_rect.update_text("Please enter your confirm password.")
            elif self.email_rect.text == "Email" or self.email_rect.text == '':
                self.announcement_rect.update_text("Please enter your email: 'example@gmail.com'.")
            elif self.conpasword_rect.text != self.password_rect.text :
                self.username_rect.update_text('Username')
                self.password_rect.update_text('Password')
                self.conpasword_rect.update_text('Confirm password')
                self.email_rect.update_text('Email')
                self.announcement_rect.update_text('Passwords do not match. Please try again.')
            elif self.gmail.is_gmail_address(self.email_rect.text) == False:
                self.announcement_rect.update_text("Email address is not valid.Please enter your email again.")
            if not self.samename:
                doc = Document(self.docx_filename)
                for table in doc.tables:
                    for row in table.rows:
                        stored_username = row.cells[1].text
                        if self.username_rect.text == stored_username:
                            self.announcement_rect.update_text("The account already exists, please give different username.")
                            self.samename = 1
            if not self.samename:
                if self.gmail.is_gmail_address(self.email_rect.text) == True:
                    if self.gmail.send_email(self.email_rect.text):
                        self.announcement_rect.update_text("Check your gmail, then enter the Verification Code.")
                        self.verify_check = True
                        self.SignUp_button.name = self.Game_font_button.render('RESEND', True, (0, 0, 0))
                    else:
                        self.announcement_rect.update_text("Email address is not valid.Please enter your email again.")
            self.samename = 0
 
    def Login(self):
        if self.mode == 'SIGN UP':
            self.username_rect.update_text('Username')
            self.password_rect.update_text('Password')
            self.conpasword_rect.update_text('Confirm password')
            self.email_rect.update_text('Email')
            self.announcement = ''
            self.announcement_rect.update_text(self.announcement)
            self.mode = 'LOGIN'
            self.check = 'LOGIN'
        elif self.mode == 'LOGIN':
            self.check = 'LOGIN'
            if self.username_rect.text == "Username" or self.username_rect.text == '':
                self.announcement_rect.update_text("Please enter your username.")
            elif self.password_rect.text == "Password" or self.password_rect.text == '':
                self.announcement_rect.update_text("Please enter your password.")
            elif self.check_login(self.username_rect.text, self.password_rect.text):
                try:
                    doc = Document(self.docx_filename)
                    for table in doc.tables:
                        for row in table.rows:
                            stored_username = row.cells[1].text
                            myID = row.cells[0].text
                            if self.username_rect.text == stored_username:
                                return myID
                    return False
                except Exception as e:
                    return False


    def Send(self):
        if self.gmail.verification_code == self.verify_rect.text:
            self.gmail_codecheck = True
        else:
            self.verify_rect.update_text('')
            self.announcement_rect.update_text("Sorry, the verification code you entered is incorrect.")
        if self.gmail_codecheck:
            self.win_money = 0
            self.lose_money = 0
            if self.save_account(self.username_rect.text, self.password_rect.text, self.money, self.match, self.win, self.lose, self.win_money, self.lose_money):
                return True
            else:
                self.announcement_rect.update_text(f"The account '{self.username_rect.text}' already exists. Cannot add.")

    def Run(self):
        LoginGame = 0
        while not LoginGame:
            for event in pygame.event.get():
                if event.type == pygame.QUIT:
                    pygame.quit()
                    sys.exit()
                if event.type == pygame.MOUSEBUTTONDOWN:
                    x, y = pygame.mouse.get_pos()
                    if self.username_rect.rect.collidepoint(x, y):
                        if self.password_rect.text == '':
                            self.password_rect.update_text("Password")
                        if self.conpasword_rect.text == '':
                            self.conpasword_rect.update_text("Confirm password")
                        if self.email_rect.text == '':
                            self.email_rect.update_text("Email")
                        if self.verify_rect.text == '':
                            self.verify_rect.update_text("Verity")
                        self.cursor_click = [True, False, False, False, False]
                        self.check = "Username"
                        self.username_rect.update_text('')
                        self.announcement_rect.update_text('')
                    elif self.password_rect.rect.collidepoint(x, y):
                        if self.username_rect.text == '':
                            self.username_rect.update_text("Username")
                        if self.conpasword_rect.text == '':
                            self.conpasword_rect.update_text("Confirm password")
                        if self.email_rect.text == '':
                            self.email_rect.update_text("Email")
                        if self.verify_rect.text == '':
                            self.verify_rect.update_text("Verity")
                        self.cursor_click = [False, True, False, False, False]
                        self.check = "Password"
                        self.password_rect.update_text('')
                        self.announcement_rect.update_text('')
                    elif self.conpasword_rect.rect.collidepoint(x, y):
                        if self.username_rect.text == '':
                            self.username_rect.update_text("Username")
                        if self.password_rect.text == '':
                            self.password_rect.update_text("Password")
                        if self.email_rect.text == '':
                            self.email_rect.update_text("Email")
                        if self.verify_rect.text == '':
                            self.verify_rect.update_text("Verity")
                        self.cursor_click = [False, False, True, False, False]
                        self.check = "Confirm password"
                        self.conpasword_rect.update_text('')
                        self.announcement_rect.update_text('')
                    elif self.email_rect.rect.collidepoint(x, y):
                        if self.username_rect.text == '':
                            self.username_rect.update_text("Username")
                        if self.password_rect.text == '':
                            self.password_rect.update_text("Password")
                        if self.conpasword_rect.text == '':
                            self.conpasword_rect.update_text("Confirm password")
                        if self.verify_rect.text == '':
                            self.verify_rect.update_text("Verity")
                        self.cursor_click = [False, False, False, True, False]
                        self.check = "Email"
                        self.email_rect.update_text('')    
                        self.announcement_rect.update_text('')
                    elif self.SignUp_button.rect.collidepoint(x, y):
                        self.announcement_rect.update_text('')
                        self.cursor_click = [False, False, False, False, False]
                        self.SignUp()
                    elif self.Login_button.rect.collidepoint(x, y):
                        self.announcement_rect.update_text('')
                        self.cursor_click = [False, False, False, False, False]
                        self.ID_Login[0] = self.Login()
                        if self.ID_Login[0]:
                            LoginGame = 1
                            break
                        else:
                            self.username_rect.update_text('Username')
                            self.password_rect.update_text('Password')
                            self.announcement_rect.update_text('Incorrect username or password.')
                    elif self.verify_check and self.verify_rect.rect.collidepoint(x, y):
                        if self.username_rect.text == '':
                            self.username_rect.update_text("Username")
                        if self.password_rect.text == '':
                            self.password_rect.update_text("Password")
                        if self.conpasword_rect.text == '':
                            self.conpasword_rect.update_text("Confirm password")
                        if self.email_rect.text == '':
                            self.email_rect.update_text("Email")
                        self.cursor_click = [False, False, False, False, True]
                        self.check = "Verity"
                        self.verify_rect.update_text('')
                        self.announcement_rect.update_text('')
                    elif self.verify_check and self.send_rect.rect.collidepoint(x, y):
                        self.cursor_click = [False, False, False, False, False]
                        self.button_click_send = True
                        self.button_click_time = pygame.time.get_ticks()
                        if self.Send():
                            time.sleep(2)
                            self.verify_check = False
                            self.SignUp_button.name = self.Game_font_button.render('SIGN UP', True, (0, 0, 0))
                            Aready.Aready()
                            TakePhotos.TakePhotos(Login_Game.Count_id())
                            self.ID_Login[0] = Login_Game.Count_id()
                            LoginGame = 1
                            break
                    elif self.Face_ID_1_rect.collidepoint(x, y):
                            self.Face_ID_1_click = True
                            self.Face_ID_click_time = pygame.time.get_ticks()
                            self.ID_Login[0] = FaceID.FaceID(FaceID.classNames)
                            if  self.ID_Login[0]:
                                LoginGame = 1
                                break
                    else:
                        if self.username_rect.text == '':
                            self.username_rect.update_text("Username")
                        if self.password_rect.text == '':
                            self.password_rect.update_text("Password")
                        if self.conpasword_rect.text == '':
                            self.conpasword_rect.update_text("Confirm password")
                        if self.email_rect.text == '':
                            self.email_rect.update_text("Email")
                        if self.verify_rect.text == '':
                            self.verify_rect.update_text("Verity")
                        self.cursor_click = [False, False, False, False, False]
                if event.type == pygame.KEYDOWN:
                    if self.check == "Username":
                        if event.key == K_BACKSPACE:
                            self.username_rect.update_text(self.username_rect.text[:-1])
                        else:
                            if len(self.username_rect.text) < 25:
                                self.username_rect.update_text(self.username_rect.text + event.unicode)
                    elif self.check == "Password":
                        if self.username_rect.text == '':
                            self.username_rect.update_text("Username")
                        if event.key == K_BACKSPACE:
                            self.password_rect.update_text(self.password_rect.text[:-1])
                        else:
                            if len(self.password_rect.text) < 25:
                                self.password_rect.update_text(self.password_rect.text + event.unicode)
                    elif self.mode == 'SIGN UP' and self.check == "Confirm password":
                        if self.password_rect.text == '':
                            self.password_rect.update_text("Password")
                        if event.key == K_BACKSPACE:
                            self.conpasword_rect.update_text(self.conpasword_rect.text[:-1])
                        else:
                            if len(self.conpasword_rect.text) < 25:
                                self.conpasword_rect.update_text(self.conpasword_rect.text + event.unicode)
                    elif self.mode == 'SIGN UP' and self.check == "Email":
                        if self.conpasword_rect.text == '':
                            self.conpasword_rect.update_text("Confirm password")
                        if event.key == K_BACKSPACE:
                            self.email_rect.update_text(self.email_rect.text[:-1])
                        else:
                            if len(self.email_rect.text) < 40:
                                self.email_rect.update_text(self.email_rect.text + event.unicode)
                    elif self.mode == 'SIGN UP' and self.check == "Verity":
                        if self.email_rect.text == '':
                            self.email_rect.update_text("Email")    
                        if event.key == K_BACKSPACE:
                            self.verify_rect.update_text(self.verify_rect.text[:-1])
                        else:
                            if len(self.verify_rect.text) < 15:
                                self.verify_rect.update_text(self.verify_rect.text + event.unicode)
            if(LoginGame):
                break
            self.screen.fill((255, 255, 255))
            self.screen.blit(self.background, (0, 0))
            pygame.draw.rect(self.screen, (220, 240, 230), (self.screenWidth // 3 - 50, self.screenHeight // 3 - 100, 500, 480), border_radius = 20)
            self.Setup.draw()
            if self.mode == 'LOGIN':
                self.Setup.text = self.mode
            elif self.mode == 'SIGN UP':
                self.Setup.text = self.mode
            self.username_rect.draw()
            if self.cursor_click[0]:
                self.username_rect.update_cursor()
            self.password_rect.draw()
            if self.cursor_click[1]:
                self.password_rect.update_cursor()
            if self.mode == 'SIGN UP':
                self.conpasword_rect.draw()
                if self.cursor_click[2]:
                    self.conpasword_rect.update_cursor()
                self.email_rect.draw()
                if self.cursor_click[3]:
                    self.email_rect.update_cursor()
            if self.mode == 'SIGN UP':
                self.SignUp_button.draw(10, (24, 119, 242))
            else:
                self.SignUp_button.draw(10, (240, 240, 240))
            if self.mode == 'LOGIN':
                self.Login_button.draw(10, (24, 119, 242))
            else:
                self.Login_button.draw(10, (240, 240, 240))
            if self.Face_ID_1_click and pygame.time.get_ticks() - self.Face_ID_click_time >= 400:
                self.Face_ID_1_click = False
            if self.Face_ID_1_click:
                self.screen.blit(self.Face_ID_2, self.Face_ID_1_rect.topleft)
            else:
                self.screen.blit(self.Face_ID_1, self.Face_ID_1_rect.topleft)
            self.announcement_rect.draw()
            if self.mode == 'SIGN UP' and self.verify_check:
                self.verify_rect.draw()
                if self.cursor_click[4]:
                    self.verify_rect.update_cursor()
                if self.button_click_send and pygame.time.get_ticks() - self.button_click_time >= 400:
                    self.button_click_send = False
                if self.button_click_send:
                    self.send_rect.draw(0, (24, 119, 242))
                else:
                    self.send_rect.draw(0, (240, 240, 240))
            pygame.display.update()
            self.Clock.tick(60)

Login_Game = Login_Game()
Login_Game.Run()

def add_mini_money(id_user, moremoney):
	docx_filename = "Accounts.docx"
	doc = Document(docx_filename)
	for table in doc.tables:
		for row in table.rows:
			if row.cells[0].text == id_user:
				row.cells[3].text = str(moremoney + int(row.cells[3].text))
				doc.save(docx_filename)
class MiniPlayer(pygame.sprite.Sprite):
	def __init__(self):
		super().__init__()
		player_walk_1 = pygame.image.load('graphics/player/player_walk_1.png').convert_alpha()
		player_walk_2 = pygame.image.load('graphics/player/player_walk_2.png').convert_alpha()
		self.player_walk = [player_walk_1,player_walk_2]
		self.player_index = 0
		self.player_jump = pygame.image.load('graphics/player/jump.png').convert_alpha()

		self.image = self.player_walk[self.player_index]
		self.rect = self.image.get_rect(midbottom = (80,600))
		self.gravity = 0

		self.jump_sound = pygame.mixer.Sound('audio/jump.mp3')
		self.jump_sound.set_volume(0.1)

	def player_input(self):
		keys = pygame.key.get_pressed()
		if keys[pygame.K_SPACE] and self.rect.bottom >= 600:
			self.gravity = -20
			self.jump_sound.play()

	def apply_gravity(self):
		self.gravity += 1
		self.rect.y += self.gravity
		if self.rect.bottom >= 600:
			self.rect.bottom = 600

	def animation_state(self):
		if self.rect.bottom < 600: 
			self.image = self.player_jump
		else:
			self.player_index += 0.1
			if self.player_index >= len(self.player_walk):self.player_index = 0
			self.image = self.player_walk[int(self.player_index)]
	
	def update(self):
		self.player_input()
		self.apply_gravity()
		self.animation_state()

class Obstacle(pygame.sprite.Sprite):
	def __init__(self,type,game):
		super().__init__()
		self.game = game
		if type == 'fly':
			fly_1 = pygame.image.load('graphics/fly/fly1.png').convert_alpha()
			fly_2 = pygame.image.load('graphics/fly/fly2.png').convert_alpha()
			self.frames = [fly_1,fly_2]
			y_pos = 450
		else:
			snail_1 = pygame.image.load('graphics/snail/snail1.png').convert_alpha()
			snail_2 = pygame.image.load('graphics/snail/snail2.png').convert_alpha()
			snail_3 = pygame.image.load('graphics/snail/snail3.png').convert_alpha()
			snail_4 = pygame.image.load('graphics/snail/snail4.png').convert_alpha()
			self.frames = [snail_1,snail_2,snail_3,snail_4]
			y_pos  = 600

		self.animation_index = 0
		self.image = self.frames[self.animation_index]
		self.rect = self.image.get_rect(midbottom = (random.randint(self.game.width + 100,self.game.width + 300),y_pos))

	def animation_state(self):
		self.animation_index += 0.1 
		if self.animation_index >= len(self.frames): self.animation_index = 0
		self.image = self.frames[int(self.animation_index)]

	def update(self):
		self.animation_state()
		self.rect.x -= 6
		self.destroy()

	def destroy(self):
		if self.rect.x <= -100: 
			self.kill()

class MiniGame():
	def display_score(self):
		self.current_time = int(pygame.time.get_ticks() / 1000) - self.start_time
		self.score_surf = self.test_font.render(f'Score: {self.current_time}',False,(64,64,64))
		self.score_rect = self.score_surf.get_rect(center = (self.width/2,self.height*0.2))
		self.display.blit(self.score_surf,self.score_rect)
		return self.current_time

	def collision_sprite(self):
		if pygame.sprite.spritecollide(self.player.sprite,self.obstacle_group,False):
			self.obstacle_group.empty()
			return False
		else: return True

	def __init__(self):
		pygame.init()
		pygame.display.set_caption('Runner')
		self.clock = pygame.time.Clock()
		self.width = 1400
		self.height = 700
		self.screen = pygame.display.set_mode((self.width,self.height))
		self.display = pygame.Surface((self.width,self.height))
		self.test_font = pygame.font.Font('font/Pixeltype.ttf', 50)
		self.game_active = False
		self.start_time = 0
		self.score = 0
		self.cursor_point = pygame.image.load('data/cursor/cursor_point2.png').convert_alpha()
		self.cursor_click = pygame.image.load('data/cursor/cursor_click2.png').convert_alpha()
		pygame.mixer.music.load('audio/music.wav')
		pygame.mixer.music.play(-1, 0, 3000)
		pygame.mixer.music.set_volume(0.1)
		self.menu_out_fx = pygame.mixer.Sound('data/sounds/menu_out.wav')
		self.menu_out_fx.set_volume(0.2)
		self.mouse_pos = pygame.mouse.get_pos()
		self.display = pygame.Surface(((self.width, self.height)))
		self.base_color = "WHITE"
		self.hovering_color = "#e2446c"
		self.LightPixel_font = pygame.font.Font('data/font/LightPixel.ttf', 30)
		self.button_image = pygame.image.load('data/small.png').convert_alpha()
		self.RETURNBUTTON = Button(image = self.button_image, pos = (self.width * 0.15, self.height * 0.5), text_input = "<Return>",
										font = self.LightPixel_font, base_color = self.base_color, hovering_color = self.hovering_color)
		self.button_pressed = False
		self.player_ID = Login_Game.get_ID()
		self.trigger = False
		self.screen_trigger = False
		#Groups
		self.player = pygame.sprite.GroupSingle()
		self.player.add(MiniPlayer())

		self.obstacle_group = pygame.sprite.Group()

		self.sky_surface = pygame.transform.scale(pygame.image.load('graphics/Sky.png'), (self.width, self.height)).convert_alpha()
		self.ground_surface = pygame.transform.scale(pygame.image.load('graphics/ground.png'), (self.width, self.height)).convert_alpha()

		# Intro screen
		self.player_stand = pygame.image.load('graphics/player/player_stand.png').convert_alpha()
		self.player_stand = pygame.transform.rotozoom(self.player_stand,0,2)
		self.player_stand_rect = self.player_stand.get_rect(center = (self.width/2,self.height/2))

		self.game_name = self.test_font.render('Pixel Runner',False,(111,196,169))
		self.game_name_rect = self.game_name.get_rect(center = (self.width/2,self.height*0.3))

		self.game_message = self.test_font.render('Press space to run',False,(111,196,169))
		self.game_message_rect = self.game_message.get_rect(center = (self.width/2,self.height*0.7))

		# Timer 
		self.obstacle_timer = pygame.USEREVENT + 1
		pygame.time.set_timer(self.obstacle_timer,1500)
		self.minigame_state = True
	
	def minigame(self, width, height, fullscreen):
		while self.minigame_state:
			for event in pygame.event.get():
				if event.type == pygame.QUIT:
					pygame.quit()
					exit()

				if self.game_active:
					if event.type == self.obstacle_timer:
						self.obstacle_group.add(Obstacle(random.choice(['fly','snail','snail','snail']), self))
				else:
					if event.type == pygame.KEYDOWN and event.key == pygame.K_SPACE and self.trigger == False:
						self.game_active = True
						self.start_time = int(pygame.time.get_ticks() / 1000)
			
			if fullscreen and self.screen_trigger == False:
				self.width = width
				self.height = height
				self.screen = pygame.display.set_mode((self.width,self.height), pygame.FULLSCREEN)
				self.player_stand_rect = self.player_stand.get_rect(center = (self.width/2,self.height/2))
				self.sky_surface = pygame.transform.scale(pygame.image.load('graphics/Sky.png'), (self.width, self.height)).convert_alpha()
				self.ground_surface = pygame.transform.scale(pygame.image.load('graphics/ground.png'), (self.width, self.height)).convert_alpha()
				self.game_name_rect = self.game_name.get_rect(center = (self.width/2,self.height*0.3))
				self.game_message_rect = self.game_message.get_rect(center = (self.width/2,self.height*0.7))
				self.screen_trigger = True
			elif fullscreen == False and self.screen_trigger == False:
				self.width = width
				self.height = height
				self.screen = pygame.display.set_mode((self.width,self.height))
				self.player_stand_rect = self.player_stand.get_rect(center = (self.width/2,self.height/2))
				self.sky_surface = pygame.transform.scale(pygame.image.load('graphics/Sky.png'), (self.width, self.height)).convert_alpha()
				self.ground_surface = pygame.transform.scale(pygame.image.load('graphics/ground.png'), (self.width, self.height)).convert_alpha()
				self.game_name_rect = self.game_name.get_rect(center = (self.width/2,self.height*0.3))
				self.game_message_rect = self.game_message.get_rect(center = (self.width/2,self.height*0.7))
				self.screen_trigger = True
			self.mouse_pos = pygame.mouse.get_pos()
			self.display = pygame.Surface(((self.width, self.height)))

			if self.game_active:
				self.display.blit(self.sky_surface,(0,0))
				self.display.blit(self.ground_surface,(0,600))
				self.score = self.display_score()
				
				self.player.draw(self.display)
				self.player.update()

				self.obstacle_group.draw(self.display)
				self.obstacle_group.update()

				self.game_active = self.collision_sprite()
				if self.score >= 20:
					self.game_active = False
				
			else:
				self.display.fill((94,129,162))
				self.display.blit(self.player_stand,self.player_stand_rect)

				score_message = self.test_font.render(f'Your earned: {self.score*5}',False,(111,196,169))
				score_message_rect = score_message.get_rect(center = (self.width/2,self.height*0.7))
				self.display.blit(self.game_name,self.game_name_rect)
				if self.score == 0: 
					self.display.blit(self.game_message,self.game_message_rect)
				else:
					self.trigger = True
					self.display.blit(score_message,score_message_rect)
				
				if self.score * 5 >= 100:
					self.score = 20
     
				#Return button
				for button in [self.RETURNBUTTON]:
					button.changeColor(self.mouse_pos)
					button.update(self.display)
					
				if self.RETURNBUTTON.checkForInput(self.mouse_pos):
					if pygame.mouse.get_pressed()[0] == 1:
						self.button_pressed = True
					elif pygame.mouse.get_pressed()[0] == 0 and self.button_pressed == True:
						self.menu_out_fx.play()
						add_mini_money(self.player_ID, self.score*5)
						pygame.mixer.music.stop()
						self.minigame_state = False
						self.button_pressed = False
						Game().run(self.width, self.height, fullscreen)

			#Mouse cursor
			pygame.mouse.set_visible(False)
			if pygame.mouse.get_pressed()[0]:
				self.display.blit(self.cursor_click, (self.mouse_pos[0] - 10, self.mouse_pos[1] - 10))
			else:
				self.display.blit(self.cursor_point, (self.mouse_pos[0] - 10, self.mouse_pos[1] - 10))
            
			self.screen.blit(pygame.transform.scale(self.display, self.screen.get_size()), (0,0))        
			pygame.display.update()
			self.clock.tick(60)
   
class Game:
    def __init__(self):
        pygame.mixer.pre_init(44100, 16, 2, 4096)
        mixer.init()
        pygame.init()
        #Sreen and surface...
        self.screen_size = [1400, 700]
        self.screen = pygame.display.set_mode((self.screen_size[0], self.screen_size[1]))
        self.width = self.screen.get_width()
        self.height = self.screen.get_height()
        self.display = pygame.Surface(((self.screen.get_width(), self.screen.get_height())))
        self.cursor_point = pygame.image.load('data/cursor/cursor_point2.png').convert_alpha()
        self.cursor_click = pygame.image.load('data/cursor/cursor_click2.png').convert_alpha()
        pygame.display.set_caption('Gearheads&Gamblers')
        self.fullscreen = False

        #Clock, font... 
        self.clock = pygame.time.Clock()
        self.LightPixel_font = pygame.font.Font('data/font/LightPixel.ttf', 30)
        self.SmallerLightPixel_font = pygame.font.Font('data/font/LightPixel.ttf', 20)
        self.BiggerLightPixel_font = pygame.font.Font('data/font/LightPixel.ttf', 40)
        self.SoftnBig_font = pygame.font.Font('data/font/Pixeltype.ttf', 50)
        self.Arcade_font = pygame.font.Font('data/font/Arcade.ttf', 60)
        self.DIRECTION_font = pygame.font.Font('data/font/LightPixel.ttf', 100)
        
        #User's variable
        self.ID = Login_Game.get_ID()
        self.player_name = get_information(self.ID)[1]
        self.money = get_information(self.ID)[3]
        self.bet = get_information(self.ID)[3]
        self.total_game = get_information(self.ID)[4]
        self.win_game = get_information(self.ID)[5]
        self.lose_game = get_information(self.ID)[6]
        self.win_money = get_information(self.ID)[7]
        self.lose_money = get_information(self.ID)[8]
        
        #Game variable
        self.player_set = 1
        self.player_index = 1
        self.player_status = 1
        self.race_started = False
        self.base_color = "WHITE"
        self.current_language = 'en'
        self.trigger = False
        self.buff_type = 'backward'
        #Map variable
        self.map_index = 1
        if self.current_language == 'vn':
            self.map_size = "NHO"
        else:
            self.map_size = "SMALL"

        #Assets
        self.assets = {
            'loadgame' : load_image('data/graphics/interface/BACKGROUND_2.png'),
            'background' : load_image('data/graphics/background/background.png'),
            'players' : load_images(f'data/player/{self.player_set}/{self.player_index}/{self.player_status}/'),
            'map_preview' : load_images("data/map/preview"),
            'map' : load_image(f'data/map/{self.map_index}/{self.map_size}.png'),
            'cars' : load_images(f'data/graphics/car/{self.player_set}/{self.player_index}'),
            'board' : load_image('data/graphics/background/board.png'),
            'ranking' : load_image('data/graphics/background/ranking.png'),
            'minigame1' : load_image('data/graphics/background/minigame1.png'),
            'minigame2' : load_image('data/graphics/background/minigame2.png'),
            'how_to_play' : load_images('data/graphics/how_to_play'),
        }
        
        self.player_names = {
            '1' : ['Dai Hiep', 'Dinh Hieu', 'Quang Duy', 'Manh Hien', 'Trong Hieu'],
            '2' : ['Chicken','Cow','Frog','Supeerduck', 'Worm'],
            '3' : ['Hoodiewoodie','Knight','Samurai','Satyr', 'Wizard'],
            '4' : ['Burger','Coin', 'Fire','Slime', 'Thunder'],
            '5' : ['Death','Eyes','Magician','Warrior', 'Witch'],
        }
        
        self.language_resources = {
            'en' : {
                'Press SPACE' : '>Press SPACE to continue<',
                'Hi' : 'Hi',
                'Quit' : '<QUIT>',
                'Language' : '<LANGUAGE>',
                'Music' : '<MUSIC>',
                'History' : '<HISTORY>',
                'Minigame' : '<MINIGAME>',
                'Map size' : 'MAP SIZE',
                'Credits' : '<CREDITS>',
                'Size' : '<SIZE>',
                'Select map' : '>SELECT THIS MAP<',
                'Small' : 'SMALL',
                'Medium' : 'MEDIUM',
                'Large' : 'LARGE',
                'Money' : 'Money',
                'Back' : '<BACK',
                'Previous' : '<PREV',
                'Next' : 'NEXT>',
                'Bet money' : 'BET MONEY',
                'Choose your player set' : '>CHOOSE YOUR PLAYER SET<',
                'Choose your player' : '>CHOOSE YOUR PLAYER<',
                'Confirm' : '<CONFIRM>',
                'Rank' : 'Rank',
                'Leaderboard' : 'LEADERBOARD',
                'Result' : 'RESULT',
                'Your rank' : 'Your rank',
                'You earned' : 'You earned',
                'You lost' : 'You lost',
                'Total money' : 'Total money',
                'Top' : 'Top',
                'Project end term' : 'PROJECT END-TERM',
                'Project manager' : 'PROJECT MANAGER:',
                'Bussiness analyst' : 'BUSINESS ANALYST:',
                'Developers' : 'DEVELOPERS:',
                'Return' : '<RETURN>',
                'Minigame limit' : '! LIMIT FOR MINIGAME MONEY: 100 !',
                'Total game' : 'Total game',
                'Win game' : 'Win game',
                'Lose game' : 'Lose game',
                'Total money win' : 'Total money win on gambling:',
                'Total money lost' : 'Total money lost on gambling:',
                'How to play' : '<How to play>',
                'You' : ' (You)',
                'minigame1' : 'MINIGAME 1',
                'minigame2' : 'MINIGAME 2',
                'Not enough money' : 'Not enough money to bet',
            },

            'vn': {
                'Press SPACE' : '>Nhan SPACE de tiep tuc<',
                'Hi' : 'Chao',
                'Quit' : '<Thoat>',
                'Language' : '<NGON NGU>',
                'Music' : '<AM NHAC>',
                'History' : '<LICH SU>',
                'Select map' : '>CHON BAN DO<',
                'Map size' : 'KICH THUOC BAN DO',
                'Small' : 'NHO',
                'Medium' : 'VUA',
                'Large' : 'LON',
                'Money' : 'Tien',
                'Minigame' : '<MINIGAME>',
                'Size' : '<KICH CO>',
                'Credits' : '<THONG TIN>',
                'Back' : '<QUAY LAI',
                'Previous' : '<TRUOC',
                'Next' : 'TIEP>',
                'Bet money' : 'DAT CUOC',
                'Choose your player set' : '>CHON BO NGUOI CHOI<',
                'Confirm' : '<XAC NHAN>',
                'Choose your player' : 'CHON NGUOI CHOI',
                'Rank' : 'Xep hang',
                'Leaderboard' : 'BANG XEP HANG',
                'Result' : 'Ket qua',
                'Your rank' : 'Ban xep hang',
                'You earned' : 'Ban da kiem duoc',
                'You lost' : 'Ban da mat',
                'Total money' : 'Tong so tien',
                'Top' : 'Hang',
                'Project end term' : 'DO AN CUOI KY',
                'Project manager' : 'TRUONG NHOM:',
                'Bussiness analyst' : 'PHAN TICH THIET KE:',
                'Developers' : 'LAP TRINH VIEN:',
                'Return' : '<QUAY LAI>',
                'Minigame limit' : '! GIOI HAN TIEN TU MINIGAME: 100 !',
                'Total game' : 'Tong so tran',
                'Win game' : 'So tran thang',
                'Lose game' : 'So tran thua',
                'Total money win' : 'Tong so tien thang cuoc:',
                'Total money lost' : 'Tong so tien thua cuoc:',
                'How to play' : '<Huong dan>',
                'You' : ' (Ban)',
                'minigame1' : 'MINIGAME 1',
                'minigame2' : 'MINIGAME 2',
                'Not enough money' : 'Khong du tien de cuoc',
            }
        }
        
        #Sound
        pygame.mixer.music.load("data/sounds/theme.wav")
        pygame.mixer.music.queue("data/sounds/race.wav")
        pygame.mixer.music.set_volume(0.1)
        self.click_fx = pygame.mixer.Sound('data/sounds/click.wav')
        self.click_fx.set_volume(0.5)
        self.menu_in_fx = pygame.mixer.Sound('data/sounds/menu_in.wav')
        self.menu_in_fx.set_volume(0.5)
        self.menu_out_fx = pygame.mixer.Sound('data/sounds/menu_out.wav')
        self.menu_out_fx.set_volume(0.2)
        self.confirm1_fx = pygame.mixer.Sound('data/sounds/confirm1.wav')
        self.confirm1_fx.set_volume(0.1)
        self.confirm2_fx = pygame.mixer.Sound('data/sounds/confirm2.wav')
        self.confirm2_fx.set_volume(0.1)
        
        #Image
        self.button_image = pygame.image.load('data/small.png').convert_alpha()
        self.dialogue_image = pygame.image.load('data/dialogue.png').convert_alpha()
        self.long_button_image = pygame.image.load('data/long.png').convert_alpha()
        
        #Game control variable
        self.game_running = True
        self.game_state = 1
        self.main_menu_state = 1
        self.size_state = 0
        self.credits_state = 0
        self.minigame_state = 0
        self.history_state = 0
        self.player_state = 0
        self.map_state = 0
        self.button_pressed = False
        self.text_state = 0
        self.finish_line_x = self.width * 0.9
        self.rank = 0
        self.music_state = 1
        self.how_to_play_state = 0
        self.screen_trigger = False
        self.minigame1_state = 0
        self.minigame2_state = 0
        self.how_index = 0
        
        #Player 
        # (Set, Type, Status) #
        self.player_group = pygame.sprite.Group()
        self.player1 = self.player_group.add(Player(self,self.player_set,1,self.player_status,(self.width * 0.2, self.height / 2)))
        self.player2 = self.player_group.add(Player(self,self.player_set,2,self.player_status,(self.width * 0.35, self.height / 2)))
        self.player3 = self.player_group.add(Player(self,self.player_set,3,self.player_status,(self.width * 0.5, self.height / 2)))
        self.player4 = self.player_group.add(Player(self,self.player_set,4,self.player_status,(self.width * 0.65, self.height / 2)))
        self.player5 = self.player_group.add(Player(self,self.player_set,5,self.player_status,(self.width * 0.8, self.height / 2)))
        self.main_player = pygame.sprite.GroupSingle()
        self.player =  self.main_player.add(Player(self,self.player_set,self.player_index,self.player_status,(self.width * 0.3, self.height / 2)))
        
        #Car
        # (Set, Type, Status) #
        self.car_group = pygame.sprite.Group()  
        self.car1 = self.car_group.add(Car(self,self.player_set,1,self.player_status,(100, self.height * 0.55)))
        self.car2 = self.car_group.add(Car(self,self.player_set,2,self.player_status,(100, self.height * 0.65)))
        self.car3 = self.car_group.add(Car(self,self.player_set,3,self.player_status,(100, self.height * 0.75)))
        self.car4 = self.car_group.add(Car(self,self.player_set,4,self.player_status,(100, self.height * 0.85)))
        self.car5 = self.car_group.add(Car(self,self.player_set,5,self.player_status,(100, self.height * 0.94)))
        self.main_car = pygame.sprite.GroupSingle()
        self.car = self.main_car.add(Car(self,self.player_set,self.player_index,self.player_status,(self.width * 0.6, self.height / 2)))

        #Timer
        self.buff_group = pygame.sprite.Group()
        self.buff_timer = pygame.USEREVENT + 1
        pygame.time.set_timer(self.buff_timer, 1000)
        self.start_time = 0
    
    def scrshoot(self):
        self.current_datetime = datetime.now()
        self.date_time_str = self.current_datetime.strftime("%H-%M_%d-%m-%Y")
        self.scrshoot_file_path = f"screenshot_folder/img-{self.date_time_str}.jpg"
        self.capture_region = (0, 0, self.width, self.height)
        self.captured_surface = self.screen.subsurface(pygame.Rect(self.capture_region))
        pygame.image.save(self.captured_surface, self.scrshoot_file_path)
     
    def current(self):
        if self.game_state != 5:
            self.current_time = 0
        else:
            self.current_time = pygame.time.get_ticks() - self.start_time
        if self.current_time > 4000:
            self.current_time = 0
        return self.current_time
    
    def get_text(self, key):
        return self.language_resources[self.current_language].get(key, '')
    
    def get_screen_state(self):
        self.width_ = self.width
        self.height_ = self.height
        self.fullscreen_ = self.fullscreen
        self.stats = [self.width_, self.height_, self.fullscreen_]
        return self.stats
        
    def run(self, width = 1400, height = 700, fullscreen = False):
        while self.game_running:
            #Update variable
            if self.screen_trigger == False:
                self.width = width
                self.height = height
                if fullscreen:
                    self.screen = pygame.display.set_mode((self.width,self.height), pygame.FULLSCREEN)
                else:
                    self.screen = pygame.display.set_mode((self.width,self.height))
                self.screen_trigger = True
                
            self.width = self.screen.get_width()
            self.height = self.screen.get_height()
            self.mouse_pos = pygame.mouse.get_pos()
            self.display = pygame.Surface(((self.width, self.height)))
            
            #State 1
            if self.game_state == 1:
                button_animation(self)
                self.display.blit(pygame.transform.scale(self.assets['loadgame'], (self.width, self.height)), (0, 0))
                self.startgame_surface = self.LightPixel_font.render(self.get_text('Press SPACE'), 0, (111, 196, 169))
                self.startgame_rect = self.startgame_surface.get_rect(center = (self.display.get_width() / 2, self.display.get_height() * 0.875))
                self.display.blit(self.startgame_surface, self.startgame_rect)
                if pygame.key.get_pressed()[pygame.K_SPACE]:
                    self.confirm1_fx.play()
                    pygame.mixer.music.play(-1, 0, 3000) 
                    self.game_state = 2
            
            #State 2
            elif self.game_state == 2:
                if self.main_menu_state == 1:
                    main_menu(self, self)
                elif self.credits_state == 1:
                    credits_menu(self)
                elif self.minigame_state == 1:
                    minigame_menu(self, self)
                elif self.history_state == 1:
                    history_menu(self, self)
                elif self.how_to_play_state == 1:
                    how_to_play(self)
                elif self.minigame1_state:
                    minigame1(self, self)
                    self.money = get_information(self.ID)[3]
                elif self.minigame2_state:
                    self.game_running = False
                    MiniGame().minigame(self.width, self.height, self.fullscreen)
                                
            #State 3
            elif self.game_state == 3:
                Player.update_pos(self)
                choose_player_set_menu(self)
                
            #State 4
            elif self.game_state == 4:
                choose_player_menu(self)
                self.start_time = pygame.time.get_ticks()
                
            #State 5
            elif self.game_state == 5:
                Car.update_pos(self, self)
                game_play(self)
                
                for event in pygame.event.get():
                    if event.type == self.buff_timer and self.race_started:
                        self.buff_group.add(Buff(self))
                self.buff_group.update()
                
                if self.current() < 3000 and self.current() > 0:
                    self.count_down_text = self.BiggerLightPixel_font.render(str(3 - int(self.current() / 1000)), 0, (111, 196, 169))
                    self.count_down_rect = self.count_down_text.get_rect(center = (self.display.get_width() / 2, self.display.get_height() / 2))
                    self.display.blit(self.count_down_text, self.count_down_rect)
                elif self.current() >= 3000 and self.current() < 4000:
                    self.count_down_text = self.BiggerLightPixel_font.render("GO!", 0, (111, 196, 169))
                    self.count_down_rect = self.count_down_text.get_rect(center = (self.display.get_width() / 2, self.display.get_height() / 2))
                    self.display.blit(self.count_down_text, self.count_down_rect)
                if self.current() >= 3000:
                    self.game.race_started = True
                    
            #State 6
            elif self.game_state == 6:
                ranking(self)
                
            #State 7
            elif self.game_state == 7: 
                leaderboard(self, self)
                
            #State 8
            elif self.game_state == 8:
                self.scrshoot()
                self.trigger = False
                self.money = get_information(self.ID)[3]
                if self.money < 100:
                    self.bet = 0
                else:
                    self.bet = get_information(self.ID)[3]
                self.total_game = get_information(self.ID)[4]
                self.win_game = get_information(self.ID)[5]
                self.lose_game = get_information(self.ID)[6]
                self.win_money = get_information(self.ID)[7]
                self.lose_money = get_information(self.ID)[8]
                self.game_state = 2
                
            #Mouse cursor
            pygame.mouse.set_visible(False)
            if pygame.mouse.get_pressed()[0]:
                self.display.blit(self.cursor_click, (self.mouse_pos[0] - 10, self.mouse_pos[1] - 10))
            else:
                self.display.blit(self.cursor_point, (self.mouse_pos[0] - 10, self.mouse_pos[1] - 10))
            
            #Event handling
            for event in pygame.event.get():
                if event.type == pygame.QUIT:
                    pygame.quit()
                    exit()
                if event.type == pygame.KEYDOWN:
                    if event.key == pygame.K_ESCAPE:
                        pygame.quit()
                        exit()            
                        
            #Update screen
            self.screen.blit(pygame.transform.scale(self.display, self.screen.get_size()), (0,0))        
            pygame.display.update()           
            self.clock.tick(60)

Game().run()
