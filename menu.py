from docx import Document
import pygame
from button import Button, AnimatedButton, CustomButton
from utils import get_text_size, get_image_size, update_assets, count_child_folders
from car import Car
from player import Player

# Menu
def main_menu(self, game):
    #Background
    self.game = game
    self.display.blit(pygame.transform.scale(self.assets['background'], (self.width, self.height)), (0,0))
    
    #Variable
    self.mouse_pos = pygame.mouse.get_pos()
    self.base_color = "WHITE"
    self.hovering_color = "#e2446c"
    
    #Game name
    self.game_name = self.Arcade_font.render("Gearheads&Gamblers", False, self.base_color)
    self.game_name_rect = self.game_name.get_rect(center = (self.width / 2, self.height * 0.1))
    self.display.blit(self.game_name, self.game_name_rect)
    
    #Choose map and load button
    button_load(self)
    if self.MAP_SIZE.checkForInput(self.mouse_pos):
        if pygame.mouse.get_pressed()[0] == 1:
            self.button_pressed = True
        elif pygame.mouse.get_pressed()[0] == 0 and self.button_pressed == True:
            self.click_fx.play()
            if self.game.map_size == "SMALL" or self.game.map_size == "NHO":
                self.game.map_size = self.get_text('Medium')
            elif self.game.map_size == "MEDIUM" or self.game.map_size == "VUA":
                self.game.map_size = self.get_text('Large')
            elif self.game.map_size == "LARGE" or self.game.map_size == "LON":
                self.game.map_size = self.get_text('Small')
            self.map_size = self.game.map_size
            self.button_pressed = False
        
    #Custom button
    self.right_button = CustomButton(self, None, pygame.image.load('data/button/right_button_1.png'), pygame.image.load('data/button/right_button_2.png'),
                                     (self.width * 0.3, self.height * 0.525), 'choose_map_right')

    self.left_button = CustomButton(self, None, pygame.image.load('data/button/left_button_1.png'), pygame.image.load('data/button/left_button_2.png'),
                                    (self.width * 0.7, self.height * 0.525), 'choose_map_left')
    self.right_button.update()
    self.left_button.update()

    #Load map
    self.map_surface = self.assets['map_preview'][int(self.map_state) % len(self.assets['map_preview'])]
    self.map_rect = self.map_surface.get_rect(center = (self.width / 2, self.height * 0.525))
    self.display.blit(pygame.transform.rotozoom(self.map_surface, 0, 1), self.map_rect)  
    
    #---#
    #Check if start button is pressed
    if self.START.checkForInput(self.mouse_pos):
        if pygame.mouse.get_pressed()[0] == 1:
            self.button_pressed = True
        elif pygame.mouse.get_pressed()[0] == 0:
            if self.button_pressed == True:
                self.confirm2_fx.play()
                self.game_state = 3
                self.button_pressed = False

    #Change screen size
    if self.SIZE.checkForInput(self.mouse_pos):
        if pygame.mouse.get_pressed()[0] == 1:
            self.button_pressed = True
        elif pygame.mouse.get_pressed()[0] == 0 and self.button_pressed == True:
            if self.size_state == 0:
                self.screen_size = [1200, 600]
                self.screen = pygame.display.set_mode((self.screen_size[0], self.screen_size[1]))
                self.size_state = 1
            elif self.size_state == 1:
                self.screen_size = [1400, 700]
                self.screen = pygame.display.set_mode((self.screen_size[0], self.screen_size[1]))
                self.size_state = 2
            elif self.size_state == 2:
                self.screen_size = [1920, 1080]
                self.screen = pygame.display.set_mode((self.screen_size[0], self.screen_size[1]), pygame.FULLSCREEN)
                self.size_state = 0
            self.click_fx.play()
            self.button_pressed = False
        
    #Check if credits button is pressed
    if self.CREDITS.checkForInput(self.mouse_pos):
        if pygame.mouse.get_pressed()[0] == 1 and self.credits_state == 0:
            self.click_fx.play()
            self.main_menu_state = 0
            self.credits_state = 1

    #Check if minigame button is pressed
    if self.MINIGAME.checkForInput(self.mouse_pos):
        if pygame.mouse.get_pressed()[0] == 1 and self.minigame_state == 0:
            self.click_fx.play()
            self.main_menu_state = 0
            self.minigame_state = 1
            
    #Check if quit button is pressed
    if self.QUIT.checkForInput(self.mouse_pos):
        if pygame.mouse.get_pressed()[0] == 1:
            self.menu_out_fx.play()
            pygame.quit()
            exit()
    
    #Change language
    if self.LANGUAGE.checkForInput(self.mouse_pos):
        if pygame.mouse.get_pressed()[0] == 1:
            self.button_pressed = True
        elif pygame.mouse.get_pressed()[0] == 0 and self.button_pressed == True:
            self.click_fx.play()
            if self.game.current_language == 'en':
                self.game.current_language = 'vn'
            elif self.game.current_language == 'vn':
                self.game.current_language = 'en'
            if self.game.map_size == "SMALL" or self.game.map_size == "NHO":
                self.game.map_size = self.get_text('Small')
            elif self.game.map_size == "MEDIUM" or self.game.map_size == "VUA":
                self.game.map_size = self.get_text('Medium')
            elif self.game.map_size == "LARGE" or self.game.map_size == "LON":
                self.game.map_size = self.get_text('Large')
            self.button_pressed = False
     
    #View history
    if self.HISTORY.checkForInput(self.mouse_pos):
        if pygame.mouse.get_pressed()[0] == 1:
            self.button_pressed = True
        elif pygame.mouse.get_pressed()[0] == 0 and self.button_pressed == True:
            self.click_fx.play()
            self.main_menu_state = 0
            self.history_state = 1
            self.button_pressed = False
    #Music
    if self.MUSIC.checkForInput(self.mouse_pos):
        if pygame.mouse.get_pressed()[0] == 1:
            self.button_pressed = True
        elif pygame.mouse.get_pressed()[0] == 0:
            if self.button_pressed == True:
                if self.music_state == 1:
                    self.music_state = 0
                    pygame.mixer.music.pause()
                elif self.music_state == 0:
                    self.music_state = 1
                    pygame.mixer.music.unpause()
                self.button_pressed = False
    
def credits_menu(self):
    #Background
    self.display.blit(pygame.transform.scale(self.assets['background'], (self.width, self.height)), (0,0))
    
    #Text
    self.line_1 = self.LightPixel_font.render(self.get_text('Project end term'), False, "YELLOW")
    self.line_2_1 = self.LightPixel_font.render(self.get_text('Project manager'), False, "WHITE")
    self.line_2_2 = self.LightPixel_font.render("NGUYEN QUANG DUY", False, "RED")
    self.line_3_1 = self.LightPixel_font.render(self.get_text('Bussiness analyst'), False, "WHITE")
    self.line_3_2 = self.LightPixel_font.render("VAN DINH HIEU", False, "RED")
    self.line_4 = self.LightPixel_font.render(self.get_text('Developers'), False, "WHITE")
    self.line_5 = self.LightPixel_font.render("NGUYEN QUANG HUY", False, "RED")
    self.line_6 = self.LightPixel_font.render("VAN DINH HIEU", False, "RED")
    self.line_7 = self.LightPixel_font.render("LUU TRONG HIEU", False, "RED")
    self.line_8 = self.LightPixel_font.render("TRAN DAI HIEP", False, "RED")
    self.line_9 = self.LightPixel_font.render("NGUYEN MANH HIEN", False, "RED")
    
    #Create text rect
    self.line_1_rect = self.line_1.get_rect(center = (self.width/2, self.height*0.1))
    self.line_2_1_rect = self.line_2_1.get_rect(center = ((self.width/2 - get_text_size(self,self.get_text('Project manager'))[0] * 0.55), self.height*0.2))
    self.line_2_2_rect = self.line_2_2.get_rect(center = ((self.width/2 + get_text_size(self,"NGUYEN QUANG DUY")[0] * 0.55), self.height*0.2))
    self.line_3_1_rect = self.line_3_1.get_rect(center = ((self.width/2 - get_text_size(self,self.get_text('Bussiness analyst'))[0] * 0.55), self.height*0.3))
    self.line_3_2_rect = self.line_3_2.get_rect(center = ((self.width/2 + get_text_size(self,"VAN DINH HIEU")[0] * 0.55), self.height*0.3))
    self.line_4_rect = self.line_4.get_rect(center = (self.width/2, self.height*0.4))
    self.line_5_rect = self.line_5.get_rect(center = (self.width/2, self.height*0.5))
    self.line_6_rect = self.line_6.get_rect(center = (self.width/2, self.height*0.6))
    self.line_7_rect = self.line_7.get_rect(center = (self.width/2, self.height*0.7))
    self.line_8_rect = self.line_8.get_rect(center = (self.width/2, self.height*0.8))
    self.line_9_rect = self.line_9.get_rect(center = (self.width/2, self.height*0.9))
    
    #Display rect
    self.display.blit(self.line_1, self.line_1_rect)
    self.display.blit(self.line_2_1, self.line_2_1_rect)
    self.display.blit(self.line_2_2, self.line_2_2_rect)
    self.display.blit(self.line_3_1, self.line_3_1_rect)
    self.display.blit(self.line_3_2, self.line_3_2_rect)
    self.display.blit(self.line_4, self.line_4_rect)
    self.display.blit(self.line_5, self.line_5_rect)
    self.display.blit(self.line_6, self.line_6_rect)
    self.display.blit(self.line_7, self.line_7_rect)
    self.display.blit(self.line_8, self.line_8_rect)
    self.display.blit(self.line_9, self.line_9_rect)
    
    #Return button
    self.RETURNBUTTON = Button(image = self.button_image, pos = (self.width * 0.15, self.height * 0.9), text_input = "<Return>",
                               font = self.LightPixel_font, base_color = self.base_color, hovering_color = self.hovering_color)

    for button in [self.RETURNBUTTON]:
        button.changeColor(self.mouse_pos)
        button.update(self.display)
        
    if self.RETURNBUTTON.checkForInput(self.mouse_pos):
        if pygame.mouse.get_pressed()[0] == 1:
            self.button_pressed = True
        elif pygame.mouse.get_pressed()[0] == 0 and self.button_pressed == True:
            self.menu_out_fx.play()
            self.credits_state = 0
            self.main_menu_state = 1
            self.button_pressed = False

def history_menu(self, game):
    #Background
    self.display.blit(pygame.transform.scale(self.assets['board'], (self.width, self.height)), (0,0))
    
    #Variable
    self.game = game
    self.base_color = "WHITE"
    self.hovering_color = "#e2446c"
    #Text
    self.history_surf = self.BiggerLightPixel_font.render(self.get_text('History'), False, "YELLOW")
    self.total_game_surf = self.LightPixel_font.render(f"{self.get_text('Total game')}: {self.total_game}", False, "WHITE")
    self.win_game_surf = self.LightPixel_font.render(f"{self.get_text('Win game')}: {self.win_game}", False, "WHITE")
    self.lose_game_surf = self.LightPixel_font.render(f"{self.get_text('Lose game')}: {self.lose_game}", False, "WHITE")
    self.win_money_surf = self.LightPixel_font.render(f"{self.get_text('Total money win')} {self.win_money}", False, "WHITE")
    self.lose_money_surf = self.LightPixel_font.render(f"{self.get_text('Total money lost')} {self.lose_money}", False, "WHITE")
    
    #Rect
    self.history_rect = self.history_surf.get_rect(center = (self.width/2, self.height*0.2))
    self.total_game_rect = self.total_game_surf.get_rect(center = (self.width/2, self.height*0.35))
    self.win_game_rect = self.win_game_surf.get_rect(center = (self.width/2, self.height*0.45))
    self.lose_game_rect = self.lose_game_surf.get_rect(center = (self.width/2, self.height*0.55))
    self.win_money_rect = self.win_money_surf.get_rect(center = (self.width/2, self.height*0.65))
    self.lose_money_rect = self.lose_money_surf.get_rect(center = (self.width/2, self.height*0.75))
    
    #Display rect
    self.display.blit(self.history_surf, self.history_rect)
    self.display.blit(self.total_game_surf, self.total_game_rect)
    self.display.blit(self.win_game_surf, self.win_game_rect)
    self.display.blit(self.lose_game_surf, self.lose_game_rect)
    self.display.blit(self.win_money_surf, self.win_money_rect)
    self.display.blit(self.lose_money_surf, self.lose_money_rect)
    
    #Return button
    self.RETURNBUTTON = Button(image = self.button_image, pos = (self.width * 0.15, self.height * 0.85), text_input = self.get_text('Return'),
                               font = self.LightPixel_font, base_color = self.base_color, hovering_color = self.hovering_color)

    for button in [self.RETURNBUTTON]:
        button.changeColor(self.mouse_pos)
        button.update(self.display)
    
    if self.RETURNBUTTON.checkForInput(self.mouse_pos):
        if pygame.mouse.get_pressed()[0] == 1:
            self.button_pressed = True
        elif pygame.mouse.get_pressed()[0] == 0 and self.button_pressed == True:
            self.menu_out_fx.play()
            self.history_state = 0
            self.main_menu_state = 1
            self.button_pressed = False
    
def minigame_menu(self, game):
    self.game = game
    #background
    self.display.blit(pygame.transform.scale(self.assets['background'], (self.width, self.height)), (0,0))

    #Text
    self.limit_text = self.BiggerLightPixel_font.render("! LIMIT FOR MINIGAME: 100 !", False, self.base_color)
    self.limit_text_rect = self.limit_text.get_rect(center = (self.width / 2, self.height * 0.1))
    
    #Variable
    self.mouse_pos = pygame.mouse.get_pos()
    self.base_color = "WHITE"
    self.hovering_color = "#e2446c"
    
def minigame1(self, game):
    self.game = game
    #background
    self.display.blit(pygame.transform.scale(self.assets['background'], (self.width, self.height)), (0,0))

    #Text
    self.limit_text = self.BiggerLightPixel_font.render(self.get_text('Minigame limit'), False, self.base_color)
    self.limit_text_rect = self.limit_text.get_rect(center = (self.width / 2, self.height * 0.1))
    
    #Variable
    self.mouse_pos = pygame.mouse.get_pos()
    self.base_color = "WHITE"
    self.hovering_color = "#e2446c"
    
    #Create button
    if self.game.money < 100:
        self.money_text = "+1"
    else:
        self.money_text = "+0"
    self.button1 = AnimatedButton(self,self.money_text,200,200,(self.width / 2 - 100, self.height / 2 - 100),5)
    
    #Update
    self.button1.draw()
    self.display.blit(self.limit_text, self.limit_text_rect)
    
    #Button
    self.MONEY = Button(image = self.button_image, pos = ((self.width * 0.05 + get_text_size(self, f"{self.get_text('Money')} {self.money}")[0] * 0.55), self.height * 0.5), text_input = f"{self.get_text('Money')} {self.money}",
                        font = self.LightPixel_font, base_color = self.hovering_color, hovering_color = self.hovering_color)
    self.RETURNBUTTON = Button(image = self.button_image, pos = (self.width * 0.15, self.height * 0.9), text_input = self.get_text('Return'),
                               font = self.LightPixel_font, base_color = self.base_color, hovering_color = self.hovering_color)
    
    #Button update
    for button in [self.RETURNBUTTON, self.MONEY]:
        button.changeColor(self.mouse_pos)
        button.update(self.display)
    
    #Check if return button is pressed
    if self.RETURNBUTTON.checkForInput(self.mouse_pos):
        if pygame.mouse.get_pressed()[0] == 1:
            self.button_pressed = True
        elif pygame.mouse.get_pressed()[0] == 0 and self.button_pressed == True:
            self.menu_out_fx.play()
            self.minigame_state = 0
            self.main_menu_state = 1
            self.button_pressed = False

def choose_player_set_menu(self):
    # Background
    button_animation(self)
    self.display.blit(pygame.transform.scale(self.assets['background'], (self.width, self.height)), (0,0))
    self.text_surf = self.LightPixel_font.render(self.get_text('Choose your player set'), False, self.base_color)
    self.text_rect = self.text_surf.get_rect(center=(self.width / 2, self.height * 0.1))
    self.display.blit(self.text_surf, self.text_rect)

    # Button
    self.NEXT_BUTTON = Button(image=None, pos=(self.width * 0.9, self.height * 0.9), text_input= self.get_text('Next'),
                              font=self.LightPixel_font, base_color=self.base_color, hovering_color=self.hovering_color)
    self.PREV_BUTTON = Button(image=None, pos=(self.width * 0.1, self.height * 0.9), text_input= self.get_text('Previous'),
                              font=self.LightPixel_font, base_color=self.base_color, hovering_color=self.hovering_color)
    self.CONFIRM_BUTTON = Button(image=None, pos=(self.width * 0.5, self.height * 0.9), text_input=self.get_text('Confirm'),
                                 font=self.LightPixel_font, base_color=self.base_color, hovering_color=self.hovering_color)
    self.BACK_BUTTON = Button(image=None, pos=((self.width * 0.05 + get_text_size(self, f"{self.get_text('Back')}")[0] * 0.55), self.height * 0.05), text_input=self.get_text('Back'),
                              font=self.LightPixel_font, base_color=self.base_color, hovering_color=self.hovering_color)
    
    buttons = [self.NEXT_BUTTON, self.PREV_BUTTON, self.CONFIRM_BUTTON, self.BACK_BUTTON]
    for button in buttons:
        button.changeColor(self.mouse_pos)
        button.update(self.display)

    if self.BACK_BUTTON.checkForInput(self.mouse_pos) and pygame.mouse.get_pressed()[0] == 1:
        self.button_pressed = True
    elif self.BACK_BUTTON.checkForInput(self.mouse_pos) and pygame.mouse.get_pressed()[0] == 0 and self.button_pressed:
        self.menu_out_fx.play()
        self.game_state = 2
        self.button_pressed = False

    if self.NEXT_BUTTON.checkForInput(self.mouse_pos) and pygame.mouse.get_pressed()[0] == 1:
        self.button_pressed = True
    elif self.NEXT_BUTTON.checkForInput(self.mouse_pos) and pygame.mouse.get_pressed()[0] == 0 and self.button_pressed:
        self.click_fx.play()
        self.player_set = (self.player_set % 5) + 1
        self.button_pressed = False

    if self.PREV_BUTTON.checkForInput(self.mouse_pos) and pygame.mouse.get_pressed()[0] == 1:
        self.button_pressed = True
    elif self.PREV_BUTTON.checkForInput(self.mouse_pos) and pygame.mouse.get_pressed()[0] == 0 and self.button_pressed:
        self.click_fx.play()
        self.player_set = (self.player_set - 2) % 5 + 1
        self.button_pressed = False
    
    if self.CONFIRM_BUTTON.checkForInput(self.mouse_pos) and pygame.mouse.get_pressed()[0] == 1:
        self.button_pressed = True
    elif self.CONFIRM_BUTTON.checkForInput(self.mouse_pos) and pygame.mouse.get_pressed()[0] == 0 and self.button_pressed:
        self.confirm2_fx.play()
        self.game_state = 4
        self.button_pressed = False
    
    #Show and ajust the bet money
    self.bet_money_text = self.LightPixel_font.render(f"{self.get_text('Bet money')}: {self.bet}", False, self.base_color)
    self.bet_money_text_rect = self.bet_money_text.get_rect(center = (self.width / 2, self.height * 0.8))
    self.display.blit(self.bet_money_text, self.bet_money_text_rect)
    
    #Bet button
    self.bet_money_up_button = CustomButton(self, None, pygame.image.load('data/button/plus_1.png'), pygame.image.load('data/button/plus_0.png'),
                                            (self.width * 0.65, self.height * 0.8), 'bet_plus')
    self.bet_money_down_button = CustomButton(self, None, pygame.image.load('data/button/minus_1.png'), pygame.image.load('data/button/minus_0.png'),
                                              (self.width * 0.35, self.height * 0.8), 'bet_minus')
    self.bet_money_50_button = CustomButton(self, None, pygame.image.load('data/button/bet_50_1.png'), pygame.image.load('data/button/bet_50_0.png'),
                                            (self.width * 0.7, self.height * 0.8), 'bet_50')
    self.bet_money_all_button = CustomButton(self, None, pygame.image.load('data/button/bet_all_1.png'), pygame.image.load('data/button/bet_all_0.png'),
                                             (self.width * 0.75, self.height * 0.8), 'bet_all')
    self.bet_money_up_button.update()
    self.bet_money_down_button.update()
    self.bet_money_50_button.update()
    self.bet_money_all_button.update()
    
    #Display the player
    for player in self.player_group:
        player.player_set = self.player_set
        player.update_img()
        player.update()
 
def choose_player_menu(self):
    #Background
    button_animation(self)
    
    self.display.blit(pygame.transform.scale(self.assets['background'], (self.width, self.height)), (0,0))
    self.text_surf = self.LightPixel_font.render(self.get_text('Choose your player'), False, self.base_color)
    self.text_rect = self.text_surf.get_rect(center = (self.width / 2, self.height * 0.1))
    self.display.blit(self.text_surf, self.text_rect)
    
    #Button
    self.NEXT_BUTTON = Button(image=None, pos=(self.width * 0.9, self.height * 0.9), text_input= self.get_text('Next'),
                              font=self.LightPixel_font, base_color=self.base_color, hovering_color=self.hovering_color)
    self.PREV_BUTTON = Button(image=None, pos=(self.width * 0.1, self.height * 0.9), text_input= self.get_text('Previous'),
                              font=self.LightPixel_font, base_color=self.base_color, hovering_color=self.hovering_color)
    self.CONFIRM_BUTTON = Button(image=None, pos=(self.width * 0.5, self.height * 0.9), text_input=self.get_text('Confirm'),
                                 font=self.LightPixel_font, base_color=self.base_color, hovering_color=self.hovering_color)
    self.BACK_BUTTON = Button(image=None, pos=((self.width * 0.05 + get_text_size(self, f"{self.get_text('Back')}")[0] * 0.55), self.height * 0.05), text_input=self.get_text('Back'),
                              font=self.LightPixel_font, base_color=self.base_color, hovering_color=self.hovering_color)
    
    for button in [self.NEXT_BUTTON, self.PREV_BUTTON, self.CONFIRM_BUTTON, self.BACK_BUTTON]:
        button.changeColor(self.mouse_pos)
        button.update(self.display)
    
    if self.BACK_BUTTON.checkForInput(self.mouse_pos):
        if pygame.mouse.get_pressed()[0] == 1:
            self.button_pressed = True
        elif pygame.mouse.get_pressed()[0] == 0 and self.button_pressed == True:
            self.menu_out_fx.play()
            self.game_state = 3
            self.button_pressed = False

    if self.CONFIRM_BUTTON.checkForInput(self.mouse_pos):
        if pygame.mouse.get_pressed()[0] == 1:
            self.button_pressed = True
        elif pygame.mouse.get_pressed()[0] == 0:
            if self.button_pressed == True:
                self.confirm1_fx.play()
                self.game_state = 5
                pygame.mixer.music.stop()
                pygame.mixer.music.load('data/sounds/race.wav')
                pygame.mixer.music.play(-1)
                self.button_pressed = False

    #Change player & car
    for player, car in zip(self.main_player, self.main_car):
        if self.NEXT_BUTTON.checkForInput(self.mouse_pos):
            if pygame.mouse.get_pressed()[0] == 1:
                self.button_pressed = True
            elif pygame.mouse.get_pressed()[0] == 0 and self.button_pressed == True:
                self.click_fx.play()
                player.player_index += 1
                if player.player_index == 6:
                    player.player_index = 1
                self.button_pressed = False
        
        if self.PREV_BUTTON.checkForInput(self.mouse_pos):
            if pygame.mouse.get_pressed()[0] == 1:
                self.button_pressed = True
            elif pygame.mouse.get_pressed()[0] == 0 and self.button_pressed == True:
                self.click_fx.play()
                player.player_index -= 1
                if player.player_index == 0:
                    player.player_index = 5
                self.button_pressed = False
                
        #Display the player
        player.player_set = self.player_set
        player.pos = (self.width * 0.4, self.height / 2)
        player.update_img()
        player.update()

        #Display the player's car
        car.rect.x = self.width * 0.6
        car.rect.y = self.height / 2
        car.pos = (car.rect.x, car.rect.y)
        car.player_set = self.player_set
        car.player_index = player.player_index
        car.index += 0.25
        car.image = pygame.transform.scale(self.assets['cars'][int(car.index % len(self.assets['cars']))], (300, 150))
        car.rect = car.image.get_rect(center = car.pos)
        self.display.blit(car.image, car.rect)

#Main game
def game_play(self):
    Car.update_pos(self, self.game)
    #Display the map
    self.display.fill("#3A9BDC")
    self.display.blit(pygame.transform.scale(pygame.image.load(f'data/map/{self.map_index}/{self.map_size}.png'), (self.width, self.height)), (0,0))
    
    #Display the cars
    for car, player in zip(self.car_group.sprites(), self.player_group.sprites()):
        car.player_set = player.player_set
        car.player_index = player.player_index
        car.update()
    
    #Check if car reached the finish line, if yes, update the rank
    for car in self.car_group.sprites():
        if car.rect.x >= self.finish_line_x and car.reached_finish_line:
            self.rank += 1
            car.rank = self.rank
            car.reached_finish_line = False
    
    #If all cars reached the finish line, display the next button
    self.NEXT_BUTTON = Button(image = None, pos = (self.width * 0.9, self.height * 0.2), text_input = self.get_text('Next'),
                              font = self.LightPixel_font, base_color = "WHITE", hovering_color = self.hovering_color)        
    if all([car.rank for car in self.car_group.sprites()]):           
        for button in [self.NEXT_BUTTON]:
            button.changeColor(self.mouse_pos)
            button.update(self.display)
    
    #Show the rank of each car
    for car in self.car_group.sprites():
        if car.rank:
            self.rank_text = self.SmallerLightPixel_font.render(f"{self.get_text('Rank')}: {car.rank}", False, "RED")
            self.rank_text_rect = self.rank_text.get_rect(center = (car.rect.x - 100, car.rect.y + 30))
            self.display.blit(self.rank_text, self.rank_text_rect)
    
    #Check if next button is pressed
    if self.NEXT_BUTTON.checkForInput(self.mouse_pos): 
        if pygame.mouse.get_pressed()[0] == 1:
            self.button_pressed = True
        elif pygame.mouse.get_pressed()[0] == 0 and self.button_pressed == True:
            self.click_fx.play()
            pygame.mixer.music.stop()
            if all([car.rank for car in self.car_group.sprites()]): 
                for car in self.car_group:
                    for player in self.main_player:
                        if car.player_index == player.player_index:
                            print(car.rank)
                            if car.rank == 1:
                                self.money += int(self.bet)
                                pygame.mixer.music.load('data/sounds/win.wav')
                            elif car.rank == 2:
                                self.money += int(self.bet * 0.75)
                                pygame.mixer.music.load('data/sounds/win.wav')
                            elif car.rank == 3:
                                self.money -= int(self.bet * 0.33)
                                pygame.mixer.music.load('data/sounds/lose.wav')
                            elif car.rank == 4:
                                self.money -= int(self.bet * 0.5)
                                pygame.mixer.music.load('data/sounds/lose.wav')
                            elif car.rank == 5:
                                self.money -= int(self.bet)
                                pygame.mixer.music.load('data/sounds/lose.wav')
            pygame.mixer.music.play(-1)
            self.game.race_started = False
            self.game_state = 6
            self.button_pressed = False

def ranking(self):
    #Backgound
    self.display.blit(pygame.transform.scale(self.assets['ranking'], (self.width, self.height)), (0,0))
    
    #Variable
    self.base_color = "WHITE"

    #Text
    self.top1_text = self.LightPixel_font.render(f"Top 1", False, "#e65d83")
    self.top2_text = self.LightPixel_font.render(f"Top 2", False, "#63af57")
    self.top3_text = self.LightPixel_font.render(f"Top 3", False, "#3a90d0")
    self.top4_text = self.LightPixel_font.render(f"Top 4", False, self.base_color)
    self.top5_text = self.LightPixel_font.render(f"Top 5", False, self.base_color)
    
    self.top1_rect = self.top1_text.get_rect(center = (self.width*0.5, self.height * 0.485))
    self.top2_rect = self.top2_text.get_rect(center = (self.width*0.308, self.height * 0.61))
    self.top3_rect = self.top3_text.get_rect(center = (self.width*0.7, self.height * 0.61))
    self.top4_rect = self.top4_text.get_rect(center = (self.width*0.125, self.height * 0.86))
    self.top5_rect = self.top5_text.get_rect(center = (self.width*0.88, self.height * 0.86))
    
    self.display.blit(self.top1_text, self.top1_rect)
    self.display.blit(self.top2_text, self.top2_rect)
    self.display.blit(self.top3_text, self.top3_rect)
    self.display.blit(self.top4_text, self.top4_rect)
    self.display.blit(self.top5_text, self.top5_rect)
    
    # Show the image of the all players and cars matching the rank
    for car, player in zip(self.car_group.sprites(), self.player_group.sprites()):
        if car.rank == 1:
            player.player_index = car.player_index
            player.player_status = 2
            player.index += 0.1
            player.update()
            player.image = self.assets['players'][int(player.index % len(self.assets['players']))]
            player.rect = player.image.get_rect(center = player.pos)
            player.pos = (self.width * 0.5, self.height * 0.26)


        elif car.rank == 2:
            player.player_index = car.player_index
            player.player_status = 2
            player.index += 0.1
            player.update()
            player.image = self.assets['players'][int(player.index % len(self.assets['players']))]
            player.rect = player.image.get_rect(center = player.pos)
            player.pos = (self.width * 0.3, self.height * 0.4)
  
        elif car.rank == 3:
            player.player_index = car.player_index
            player.player_status = 2
            player.index += 0.1
            player.update()
            player.image = self.assets['players'][int(player.index % len(self.assets['players']))]
            player.rect = player.image.get_rect(center = player.pos)
            player.pos = (self.width * 0.7, self.height * 0.4)

        elif car.rank == 4:
            self.rank4 = car.player_index
            player.index += 0.1
            player.update()
            player.image = self.assets['players'][int(player.index % len(self.assets['players']))]
            player.rect = player.image.get_rect(center = player.pos)
            player.pos = (self.width * 0.13, self.height * 0.65)
           
        elif car.rank == 5:
            self.rank5 = car.player_index
            player.index += 0.1
            player.update()
            player.image = self.assets['players'][int(player.index % len(self.assets['players']))]
            player.rect = player.image.get_rect(center = player.pos)
            player.pos = (self.width * 0.88, self.height * 0.65)
    
    self.NEXT_BUTTON = Button(image = None, pos = (self.width * 0.85, self.height * 0.3), text_input = self.get_text('Next'),
                              font = self.LightPixel_font, base_color = self.base_color, hovering_color = self.hovering_color)
    
    for button in [self.NEXT_BUTTON]:
        button.changeColor(self.mouse_pos)
        button.update(self.display)
    
    if self.NEXT_BUTTON.checkForInput(self.mouse_pos): 
        if pygame.mouse.get_pressed()[0] == 1:
            self.button_pressed = True
        elif pygame.mouse.get_pressed()[0] == 0 and self.button_pressed == True:
            self.click_fx.play()
            self.game_state = 7
            self.button_pressed = False

def leaderboard(self, game):
    #Variable
    self.base_color = "WHITE"
    self.hovering_color = "#e2446c"
    self.game = game
    
    #Background
    self.display.blit(pygame.transform.scale(self.assets['board'], (self.width, self.height)), (0,0))
    
    #Game name
    self.game_name = self.Arcade_font.render("Gearheads&Gamblers", False, self.base_color)
    self.line1_surf = self.LightPixel_font.render(self.get_text('Leaderboard'), False, self.base_color)
    self.game_name_rect = self.game_name.get_rect(center = (self.width / 2, self.height * 0.15))
    self.line1_rect = self.line1_surf.get_rect(center = (self.width*0.3, self.height * 0.25))
    self.display.blit(self.game_name, self.game_name_rect)
    self.display.blit(self.line1_surf, self.line1_rect)

    #Define the top 5
    for car in self.car_group.sprites():
        if car.rank == 1:
            self.top1 = car.player_index
        elif car.rank == 2:
            self.top2 = car.player_index
        elif car.rank == 3:
            self.top3 = car.player_index
        elif car.rank == 4:
            self.top4 = car.player_index
        elif car.rank == 5:
            self.top5 = car.player_index
    
    #Top text
    self.top1_surf = self.LightPixel_font.render(f"{self.get_text('Top')} 1: {self.top1}", False, self.base_color)
    self.top2_surf = self.LightPixel_font.render(f"{self.get_text('Top')} 2: {self.top2}", False, self.base_color)
    self.top3_surf = self.LightPixel_font.render(f"{self.get_text('Top')} 3: {self.top3}", False, self.base_color)
    self.top4_surf = self.LightPixel_font.render(f"{self.get_text('Top')} 4: {self.top4}", False, self.base_color)
    self.top5_surf = self.LightPixel_font.render(f"{self.get_text('Top')} 5: {self.top5}", False, self.base_color)
    self.top1_rect = self.top1_surf.get_rect(center = (self.width*0.3, self.height * 0.4))
    self.top2_rect = self.top2_surf.get_rect(center = (self.width*0.3, self.height * 0.5))
    self.top3_rect = self.top3_surf.get_rect(center = (self.width*0.3, self.height * 0.6))
    self.top4_rect = self.top4_surf.get_rect(center = (self.width*0.3, self.height * 0.7))
    self.top5_rect = self.top5_surf.get_rect(center = (self.width*0.3, self.height * 0.8))
    self.display.blit(self.top1_surf, self.top1_rect)
    self.display.blit(self.top2_surf, self.top2_rect)
    self.display.blit(self.top3_surf, self.top3_rect)
    self.display.blit(self.top4_surf, self.top4_rect)
    self.display.blit(self.top5_surf, self.top5_rect)

    #Results)
    for car in self.car_group:
        for player in self.main_player:
            if car.player_index == player.player_index:
                self.result = car.rank
    self.result_surf = self.LightPixel_font.render(self.get_text('Result'), False, self.base_color)
    self.show_rank_surf = self.LightPixel_font.render(f"{self.get_text('Your rank')}: {self.result}", False, self.base_color)
    if self.trigger == False:
        if self.result == 1:
            self.money_diff = int(self.bet)
            add_money(self.game.ID, self.money_diff)
            self.money_diff_surf = self.LightPixel_font.render(f"{self.get_text('You earned')}: {self.money_diff}", False, self.base_color)
        elif self.result == 2:
            self.money_diff = int(self.bet * 0.75)
            add_money(self.game.ID, self.money_diff)
            self.money_diff_surf = self.LightPixel_font.render(f"{self.get_text('You earned')}: {self.money_diff}", False, self.base_color)
        elif self.result == 3:
            self.money_diff = int(self.bet * 0.33)
            minus_money(self.game.ID, self.money_diff)
            self.money_diff_surf = self.LightPixel_font.render(f"{self.get_text('You lost')}: {self.money_diff}", False, self.base_color)
        elif self.result == 4:
            self.money_diff = int(self.bet * 0.5)
            minus_money(self.game.ID, self.money_diff)
            self.money_diff_surf = self.LightPixel_font.render(f"{self.get_text('You lost')}: {self.money_diff}", False, self.base_color)
        elif self.result == 5:
            self.money_diff = int(self.bet)
            minus_money(self.game.ID, self.money_diff)
            self.money_diff_surf = self.LightPixel_font.render(f"{self.get_text('You lost')}: {self.money_diff}", False, self.base_color)
        self.trigger = True
        
    self.money_surf = self.LightPixel_font.render(f"{self.get_text('Total money')}: {self.money}", False, self.base_color)
    self.money_diff_rect = self.money_diff_surf.get_rect(center = (self.width*0.7, self.height * 0.55))
    self.result_rect = self.result_surf.get_rect(center = (self.width*0.7, self.height * 0.25))
    self.show_rank_rect = self.show_rank_surf.get_rect(center = (self.width*0.7, self.height * 0.45))
    self.money_rect = self.money_surf.get_rect(center = (self.width*0.7, self.height * 0.65))
    self.display.blit(self.result_surf, self.result_rect)
    self.display.blit(self.show_rank_surf, self.show_rank_rect)
    self.display.blit(self.money_surf, self.money_rect)
    self.display.blit(self.money_diff_surf, self.money_diff_rect)
    
    #Next button
    self.NEXT_BUTTON = Button(image = None, pos = (self.width * 0.85, self.height * 0.3), text_input = self.get_text('Next'),
                            font = self.BiggerLightPixel_font, base_color = "WHITE", hovering_color = self.hovering_color) 
    for button in [self.NEXT_BUTTON]:
        button.changeColor(self.mouse_pos)
        button.update(self.display)
        
    if self.NEXT_BUTTON.checkForInput(self.mouse_pos): 
        if pygame.mouse.get_pressed()[0] == 1:
            self.button_pressed = True
        elif pygame.mouse.get_pressed()[0] == 0 and self.button_pressed == True:
            self.click_fx.play()
            for car in self.car_group.sprites():
                car.reset()
            for player in self.player_group:
                player.player_status = 1
                
            if self.result == 1 or self.result == 2:
                update_win(self.game.ID)
            else:
                update_lose(self.game.ID)
            
            pygame.mixer.music.stop()
            pygame.mixer.music.load('data/sounds/theme.wav')
            pygame.mixer.music.play(-1)
            self.game_state = 8
            self.rank = 0
            self.button_pressed = False
            if self.money < 100:
                self.bet = 0
            else:
                self.bet = self.money
    return self.result

#Button    
def button_load(self):
    button_animation(self) 
    self.START = Button(image = self.long_button_image ,pos = (self.width / 2, self.height * 0.85),text_input = self.get_text('Select map'), 
                        font = self.LightPixel_font, base_color = self.base_color, hovering_color = self.hovering_color)
    self.LANGUAGE = Button(image = self.button_image, pos = ((self.width * 0.95 - get_text_size(self, self.get_text('Language'))[0] / 2), (self.height * 0.3)),text_input = self.get_text('Language'), 
                           font = self.LightPixel_font, base_color = self.base_color, hovering_color = self.hovering_color)
    self.HISTORY = Button(image = self.button_image, pos = ((self.width * 0.95 - get_text_size(self, self.get_text('History'))[0] / 2),(self.height * 0.4)),text_input = self.get_text('History'), 
                          font = self.LightPixel_font, base_color = self.base_color, hovering_color = self.hovering_color)
    self.MUSIC = Button(image = self.button_image, pos = ((self.width * 0.95 - get_text_size(self, self.get_text('Music'))[0] / 2),(self.height * 0.5)), text_input = self.get_text('Music'),
                        font = self.LightPixel_font, base_color = self.base_color, hovering_color = self.hovering_color)
    self.SIZE = Button(image = self.button_image, pos = ((self.width * 0.95 - get_text_size(self, self.get_text('Size'))[0] / 2),(self.height * 0.6)), text_input = self.get_text('Size'),
                       font = self.LightPixel_font, base_color = self.base_color, hovering_color = self.hovering_color)
    self.CREDITS = Button(image = self.button_image, pos = ((self.width * 0.95 - get_text_size(self, self.get_text('Credits'))[0] / 2),(self.height * 0.7)), text_input = self.get_text('Credits'),
                          font = self.LightPixel_font, base_color = self.base_color, hovering_color = self.hovering_color)
    self.MINIGAME = Button(image = self.button_image, pos = ((self.width * 0.95 - get_text_size(self, self.get_text('Minigame'))[0] / 2),(self.height * 0.8)), text_input = self.get_text('Minigame'),
                           font = self.LightPixel_font, base_color = self.base_color, hovering_color = self.hovering_color)
    self.MONEY = Button(image = self.button_image, pos = ((self.width * 0.05 + get_text_size(self, f"{self.get_text('Money')} {self.money}")[0] * 0.55), self.height * 0.5), text_input = f"{self.get_text('Money')} {self.money}",
                        font = self.LightPixel_font, base_color = self.hovering_color, hovering_color = self.hovering_color)
    self.NAME = Button(image = self.long_button_image , pos = ((self.width * 0.05 + get_text_size(self, f"{self.get_text('Hi')} {self.player_name}!")[0] * 0.5), self.height * 0.3), text_input = f"{self.get_text('Hi')} {self.player_name}!",
                       font = self.LightPixel_font, base_color = self.hovering_color, hovering_color = self.hovering_color)
    self.QUIT = Button(image = self.button_image, pos = ((self.width * 0.05 + get_text_size(self, f"{self.get_text('Quit')}")[0] * 0.55), self.height * 0.9), text_input = self.get_text('Quit'),
                       font = self.LightPixel_font, base_color = self.base_color, hovering_color = self.hovering_color)
    self.MAP_SIZE = Button(image = self.button_image, pos = (self.width / 2, self.height * 0.7), text_input = f"{self.get_text('Map size')}: {self.map_size}",
                           font = self.LightPixel_font, base_color = self.base_color, hovering_color = self.base_color)
    #Button update
    for button in [self.START, self.LANGUAGE, self.HISTORY, self.MUSIC, self.SIZE,
                   self.CREDITS, self.MINIGAME, self.QUIT, self.MAP_SIZE]:
        button.changeColor(self.mouse_pos)
        button.update(self.display)
    
    for button in [self.MONEY, self.NAME]:
        button.update(self.display)

#Animation
def button_animation(self):
    self.select_map_texts = [" >SELECT THIS MAP< ", "> SELECT THIS MAP <"]
    self.startgame_texts = [">Press SPACE to start<", "> Press SPACE to start <"]
    self.choose_player_set_texts = [">Choose your player set<", "> Choose your player set <"]
    self.choose_player_texts = [">Choose your player<", "> Choose your player <"]
    self.confirm_texts = ["<CONFIRM>", "< CONFIRM >"]
    self.text_state += 0.03
    self.select_map_text = self.select_map_texts[(int(self.text_state) % 2)] 
    self.startgame_text = self.startgame_texts[(int(self.text_state) % 2)] 
    self.choose_player_set_text = self.choose_player_set_texts[(int(self.text_state) % 2)]    
    self.confirm_text = self.confirm_texts[(int(self.text_state) % 2)]
    self.choose_player_text = self.choose_player_texts[(int(self.text_state) % 2)]

#Choose Player Set    
def player_load(self):
    update_assets(self, self.assets)
    self.player_state += 0.05
    self.player_surface = self.assets['players'][int(self.player_state) % len(self.assets['players'])]
    self.player_size = get_image_size(self,self.player_surface)
    self.player_rect = self.player_surface.get_rect(bottomright = (self.width / 2 - self.player_size[0], self.height / 2 - self.player_size[1]))
    self.display.blit(pygame.transform.rotozoom(self.player_surface, 0, 5), self.player_rect)

def add_money(id_user, moremoney):
    docx_filename = "Accounts.docx"
    doc = Document(docx_filename)
    for table in doc.tables:
        for row in table.rows:
            if row.cells[0].text == id_user:
                row.cells[3].text = str(moremoney + int(row.cells[3].text))
                row.cells[7].text = str(moremoney + int(row.cells[7].text))
                doc.save(docx_filename)

def minus_money(id_user, lessmoney):
    docx_filename = "Accounts.docx"
    doc = Document(docx_filename)
    for table in doc.tables:
        for row in table.rows:
            if row.cells[0].text == id_user:
                row.cells[3].text = str(int(row.cells[3].text) - lessmoney)
                row.cells[8].text = str(int(row.cells[8].text) + lessmoney)
                doc.save(docx_filename)

def update_win(id_user):
    docx_filename = "Accounts.docx"
    doc = Document(docx_filename)
    for table in doc.tables:
        for row in table.rows:
            if row.cells[0].text == id_user:
                row.cells[4].text = str(1 + int(row.cells[4].text))
                row.cells[5].text = str(1 + int(row.cells[5].text))
                doc.save(docx_filename)

def update_lose(id_user):
    docx_filename = "Accounts.docx"
    doc = Document(docx_filename)
    for table in doc.tables:
        for row in table.rows:
            if row.cells[0].text == id_user:
                row.cells[4].text = str(1 + int(row.cells[4].text))
                row.cells[6].text = str(1 + int(row.cells[6].text))
                doc.save(docx_filename)
                
language_resources = {}

def load_sound(self):
    self.click_fx = pygame.mixer.Sound('data/sound/click.wav')
    self.click_fx.set_volume(0.5)
    self.menu_in_fx = pygame.mixer.Sound('data/sound/menu_in.wav')
    self.menu_in_fx.set_volume(0.5)
    self.menu_out_fx = pygame.mixer.Sound('data/sound/menu_out.wav')
    self.menu_out_fx.set_volume(0.5)
    self.confirm1_fx = pygame.mixer.Sound('data/sound/confirm1.wav')
    self.confirm1_fx.set_volume(0.5)
    self.confirm2_fx = pygame.mixer.Sound('data/sound/confirm2.wav')
    self.confirm2_fx.set_volume(0.5)

def get_text(self, key):
    return language_resources[self.game.current_language].get(key, '')
